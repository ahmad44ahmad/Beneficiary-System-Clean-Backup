#!/usr/bin/env python
"""Build دليل-استخدام-بصيرة.docx with embedded HRSD-styled screenshots."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

import sys
SLIM = "--slim" in sys.argv
ASSETS = r"C:\dev\basira\docs\guide\assets\screenshots_opt" if SLIM else r"C:\dev\basira\docs\guide\assets\screenshots"
EXT = ".jpg" if SLIM else ".png"
DESKTOP = r"C:\Users\aass1\Desktop"
OUT_DOCX = os.path.join(DESKTOP, ("دليل-استخدام-بصيرة-2026-slim.docx" if SLIM else "دليل-استخدام-بصيرة-2026.docx"))

NAVY = "0F3144"
TEAL = "269798"
ORANGE = "F7941D"
GOLD = "FCB614"
GREEN = "2BB574"
GRAY = "7A7A7A"
WHITE = "FFFFFF"

ARABIC_FONT = "Tajawal"


def set_run_arabic(run, font=ARABIC_FONT):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    if rFonts not in rPr:
        rPr.append(rFonts)
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    cs = OxmlElement("w:cs")
    rPr.append(cs)


def set_para_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def shade_para(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_h(doc, text, level, color=NAVY, size=None, align="right", spacing=None):
    """Add Arabic heading."""
    p = doc.add_paragraph()
    set_para_rtl(p)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if spacing is not None:
        p.paragraph_format.space_before = Pt(spacing[0])
        p.paragraph_format.space_after = Pt(spacing[1])
    run = p.add_run(text)
    run.bold = True
    set_run_arabic(run)
    sizes = {1: 36, 2: 24, 3: 18, 4: 14}
    run.font.size = Pt(size or sizes.get(level, 14))
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_para(doc, text, color="333333", size=11, bold=False, italic=False, align="right", indent=None):
    p = doc.add_paragraph()
    set_para_rtl(p)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    set_run_arabic(run)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    set_para_rtl(p)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    set_run_arabic(run)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("333333")
    return p


def add_callout(doc, role_label, text, role_color=TEAL):
    """Per-role callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    # Shade
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F4F8F9")
    tcPr.append(shd)
    # Role label
    p1 = cell.paragraphs[0]
    set_para_rtl(p1)
    r1 = p1.add_run(role_label + ": ")
    r1.bold = True
    set_run_arabic(r1)
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor.from_string(role_color)
    r2 = p1.add_run(text)
    set_run_arabic(r2)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor.from_string("333333")
    # Spacing
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)


def add_image(doc, path, caption=None, width_inches=6.5):
    if SLIM:
        path = path.replace(".png", ".jpg")
    if not os.path.exists(path):
        add_para(doc, f"[تعذَّر تحميل اللقطة: {os.path.basename(path)}]", color="DC2626", italic=True, align="center")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run("الشكل: " + caption)
        set_run_arabic(cr)
        cr.font.size = Pt(10)
        cr.italic = True
        cr.font.color.rgb = RGBColor.from_string(GRAY)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_cover(doc):
    """Cover page mimicking the example PDFs."""
    # Spacing
    for _ in range(3):
        doc.add_paragraph()
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_rtl(p)
    shade_para(p, NAVY)
    r = p.add_run("                                                  ")
    set_run_arabic(r)
    p.paragraph_format.space_after = Pt(0)

    # Real cover content
    for _ in range(2):
        doc.add_paragraph()

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_rtl(p1)
    r1 = p1.add_run("الموارد البشرية والتنمية الاجتماعية")
    set_run_arabic(r1)
    r1.font.size = Pt(14)
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(NAVY)

    for _ in range(4):
        doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_rtl(p2)
    r2 = p2.add_run("دليل الاستخدام")
    set_run_arabic(r2)
    r2.font.size = Pt(56)
    r2.bold = True
    r2.font.color.rgb = RGBColor.from_string(NAVY)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_rtl(p3)
    r3 = p3.add_run("نظام بصيرة — مركز التأهيل الشامل بمنطقة الباحة")
    set_run_arabic(r3)
    r3.font.size = Pt(20)
    r3.font.color.rgb = RGBColor.from_string(TEAL)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_rtl(p4)
    r4 = p4.add_run("الإصدار التعريفي والتشغيلي الكامل  •  مايو ٢٠٢٦")
    set_run_arabic(r4)
    r4.font.size = Pt(13)
    r4.font.color.rgb = RGBColor.from_string(GRAY)

    page_break(doc)


def add_toc(doc):
    add_h(doc, "جدول المحتويات", 1, color=NAVY, size=32)
    doc.add_paragraph()
    items = [
        ("✦", "عن نظام بصيرة — الرسالة وما يُميِّزه"),
        ("✦", "مميِّزات النظام التي تَنقلهُ نقلةً نوعيَّة"),
        ("٠١", "تَسجيل الدخول وتَهيئة الحساب"),
        ("٠٢", "الشاشة الرئيسية ولوحة القيادة التنفيذية"),
        ("٠٣", "المستفيدون — استعراض البيانات وملف الكرامة"),
        ("٠٤", "الخدمات الطبية والمتابعة اليومية"),
        ("٠٥", "الخدمات الاجتماعية ومحرك التمكين"),
        ("٠٦", "الحوكمة والجودة والتميز المؤسسي"),
        ("٠٧", "العمليات اليومية"),
        ("٠٨", "الذكاء والتنبؤ"),
        ("٠٩", "التقارير والمؤشرات الاستراتيجية"),
        ("١٠", "بوصلة القيادة الاستراتيجية"),
        ("١١", "الإدارة والصلاحيات وسجلات التدقيق"),
        ("١٢", "لوحة القيادة الإشرافية الإقليمية"),
    ]
    for num, title in items:
        p = doc.add_paragraph()
        set_para_rtl(p)
        r1 = p.add_run(f"  {title}")
        set_run_arabic(r1)
        r1.font.size = Pt(13)
        r1.font.color.rgb = RGBColor.from_string("333333")
        r2 = p.add_run(f"   {num}")
        set_run_arabic(r2)
        r2.font.size = Pt(14)
        r2.bold = True
        r2.font.color.rgb = RGBColor.from_string(ORANGE if num.startswith("٠") or num.startswith("١") else TEAL)
        p.paragraph_format.space_after = Pt(8)
    page_break(doc)


def add_divider(doc, num, title):
    """Section divider page: big orange number + navy title."""
    for _ in range(8):
        doc.add_paragraph()
    p1 = doc.add_paragraph()
    set_para_rtl(p1)
    r1 = p1.add_run(num)
    set_run_arabic(r1)
    r1.font.size = Pt(96)
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(ORANGE)

    p2 = doc.add_paragraph()
    set_para_rtl(p2)
    r2 = p2.add_run(title)
    set_run_arabic(r2)
    r2.font.size = Pt(32)
    r2.bold = True
    r2.font.color.rgb = RGBColor.from_string(NAVY)
    page_break(doc)


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    # Set section RTL
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)

    # Default style RTL/Arabic
    style = doc.styles["Normal"]
    style.font.name = ARABIC_FONT
    style.font.size = Pt(11)

    # ---- COVER ----
    add_cover(doc)

    # ---- TOC ----
    add_toc(doc)

    # ---- ABOUT ----
    add_h(doc, "عن نظام بصيرة", 1, color=NAVY, size=32)
    add_para(doc, (
        "نظام بصيرة منظومة عمليات رقمية متكاملة، صُمِّمَت لخدمة مراكز التأهيل الشامل التابعة لقطاع التنمية الاجتماعية، "
        "وتُطبَّق ابتداءً في مركز التأهيل الشامل بمنطقة الباحة كبيئة Sandbox، تَمهيداً للتعميم على ٣٨ مركزاً في عموم المملكة."
    ), size=12)
    doc.add_paragraph()
    add_h(doc, "رسالة النظام", 3, color=TEAL)
    add_para(doc, (
        "أن يَنتقل العمل في مراكز التأهيل من نمط التَوثيق الورقي والإجراءات المتفرِّقة، إلى منظومة موحَّدة تَضع المستفيد في المركز، "
        "وتُمكِّن الكوادر من تَقديم خدمة كريمة قابلة للقياس، وتُتيح للقيادات رؤيةً شاملةً للأداء والأثر."
    ))
    add_h(doc, "ما يُميِّز بصيرة", 3, color=TEAL)
    for b in [
        "النموذج الاجتماعي للإعاقة أساس التصميم: المستفيد لا يُعَامَل كحالة، بل كإنسان تُزال عن طريقه الحواجز الاجتماعية.",
        "ملف الكرامة كركن مستقلّ يُوثِّق ما يُحبُّه المستفيد، تَفضيلاته الحسيَّة، أسلوب التواصل المفضَّل لديه، وسجلّ أعماله الكريمة.",
        "محرِّكات ذكية تَعمل على بيانات المركز محلياً (مروءة لتَوزيع الكوادر، حسّة زكية لتقييم المخاطر، بوصلة القيادة لقرارات الاستراتيجية).",
        "حوكمة مُحكَمَة: التزام كامل بالضوابط الأساسية للأمن السيبراني وسياسات الوزارة لحماية البيانات الشخصية وأخلاقيات الذكاء الاصطناعي من سدايا.",
        "أَنسنة الخدمة كقرار تَصميمي، لا كشعار: المستفيد مُستفيد لا 'مريض'، التدخُّل تَمكين لا 'علاج'.",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "القيمة المضافة للمستفيد المباشر",
                "مَلَفّ شخصي لا يُنسى، خطّة تأهيل تتطوَّر معه، أسرة على تواصل مستمر، فريق رعاية مُتسق رغم تَغيُّر الكوادر، وأثر تأهيلي مَقيس ومَعروض.",
                role_color=GREEN)
    page_break(doc)

    # ---- FEATURES ----
    add_h(doc, "مميِّزات النظام", 1, color=NAVY, size=32)
    add_para(doc, "تَمتدّ مميِّزات بصيرة عبر ثلاثة محاور:", size=12)
    doc.add_paragraph()
    add_h(doc, "المحور الأول — الإنسان أوَّلاً", 3, color=TEAL)
    for b in [
        "مَلَفّ الكرامة (Dignity Profile) كركن مستقلّ.",
        "بَوابة الأسرة لتَواصل مستمرّ بين المركز وذوي المستفيد.",
        "محرِّك التَمكين (Empowerment Engine) لتَتبُّع أهداف التأهيل الفردية.",
        "لغة موحَّدة 'مستفيد' / 'حاجز' / 'تَدخُّل'.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "المحور الثاني — الكوادر بكفاءة وعدالة", 3, color=TEAL)
    for b in [
        "محرِّك مروءة للوقاية من الاحتراق المهني وتَوزيع الكوادر بعدالة.",
        "بَوابة موحَّدة لكل التَخصُّصات.",
        "نَموذج تَسليم المناوبات الذكي (Handover) لمنع فقدان المعلومة بين الورديات.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "المحور الثالث — قيادة بِبيانات لا بانطباعات", 3, color=TEAL)
    for b in [
        "لوحة قيادة تَنفيذية فورية بمؤشِّرات حقيقية.",
        "محرِّك حسّة زكية لتَقييم المخاطر السريرية والسلوكية.",
        "نظام الإنذار المبكر للحالات والعنابر.",
        "بُوصلة القيادة لحفظ القرارات الاستراتيجية ومسارها.",
        "تَقارير العائد الاجتماعي للاستثمار (SROI) للقيادات والإدارات العامة.",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "للقارئ من القيادات",
                "المؤشِّرات لم تَعُد تُجمَع يدوياً من تقارير شهرية متأخِّرة. هي مَوصولة بالبيانات الفعلية وتُحدَّث لحظياً.",
                role_color=NAVY)
    add_callout(doc, "للقارئ من رؤساء الأقسام",
                "ما كان يُنجز سابقاً في خمسة نماذج ورقية متفرِّقة، يُنجز الآن في إجراء واحد، ويصل تلقائياً إلى لوحة القيادة دون تَكرار للإدخال.",
                role_color=TEAL)
    add_callout(doc, "للقارئ من مدير المركز",
                "القرار الذي يُتَّخذ يَوم الأحد يَبقى موثَّقاً بسياقه، يَستطيع المدير اللاحق بعد سنوات أن يَقرأ لماذا اتُّخِذَ، ولماذا اختير ما اختير من البدائل.",
                role_color=ORANGE)
    page_break(doc)

    # ---- SECTION 01 ----
    add_divider(doc, "٠١", "تَسجيل الدخول وتَهيئة الحساب")
    add_h(doc, "٠١  •  تَسجيل الدخول وتَهيئة الحساب", 2, color=NAVY)
    add_para(doc, "تَتَوفَّر في بصيرة قَنَوات دخول مُتعدِّدة، وَفقَ نوع المستخدم وموقعه:", size=12)
    add_h(doc, "١-١  صفحة الدخول الأساسية", 3, color=TEAL)
    add_para(doc, "عند الانتقال إلى الرابط الرسمي للنظام، تظهر شاشة الدخول بهويَّة الوزارة المعتمَدَة.")
    add_image(doc, os.path.join(ASSETS, "basira-00-login.png"), caption="شاشة تَسجيل الدخول")
    add_h(doc, "آلية الدخول", 4, color=NAVY)
    for i, t in enumerate([
        "يَدخل المستخدم بَيانات هويَّته الوظيفية (البريد المؤسَّسي أو الهوية الوطنية) عبر النَفاذ الوطني الموحَّد.",
        "يَستلم رمز التحقُّق على جواله الرسمي المُسَجَّل لدى إدارة الموارد البشرية.",
        "يَدخل الرمز ويَنتقل تلقائياً إلى الشاشة المناسبة لدوره الوظيفي.",
    ], 1):
        add_para(doc, f"{i}. {t}")
    add_h(doc, "نَطاق الصلاحيات", 4, color=NAVY)
    for b in [
        "التَوثُّق المتعدِّد العوامل (MFA) للأدوار العالية.",
        "تَقييد الوصول حسب الدور: مدير، إداري، طبيب، اختصاصي اجتماعي، مُشرف، ممرض، موظف عام.",
        "أمن مستوى الصف (Row-Level Security): الموظف لا يَرى من بيانات المستفيدين إلا ما يَخصّ نطاق عمله.",
        "تَدقيق آلي: كل دخول وكل تَفاعُل يُسَجَّل في سجلّ تَدقيق غير قابل للتعديل.",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "للقارئ من إدارة الإشراف الاجتماعي بفرع الوزارة",
                "يَمنحك النظام رؤيةً عابرةً لكل مستفيدٍ ضمن نطاق إشرافك، مع تتبُّع كامل لِما تَمَّ من خَدَمات في كل مركز خاضع لإشرافك.",
                role_color=GREEN)
    add_callout(doc, "للقارئ من المدير العام بفرع الوزارة",
                "الدخول يُتيح لك رؤية مُجَمَّعَة لكل المراكز والوحدات في منطقتك، مع إمكانية الاستفسار العميق عند الحاجة.",
                role_color=NAVY)
    page_break(doc)

    # ---- SECTION 02 ----
    add_divider(doc, "٠٢", "الشاشة الرئيسية ولوحة القيادة التنفيذية")
    add_h(doc, "٠٢  •  الشاشة الرئيسية ولوحة القيادة التنفيذية", 2, color=NAVY)
    add_para(doc, "بعد تَسجيل الدخول، تَنفتح الشاشة الرئيسية. تَتألَّف من ثلاثة أقسام:", size=12)
    for i, t in enumerate([
        "الترويسة (أعلى الشاشة): اسم المركز، شعار الوزارة، اسم المستخدم ودوره، إشعارات اللحظة، وصلاحية تَسجيل الخروج.",
        "الشريط الجانبي (يَمين الشاشة): تسعة أقسام رئيسية تَفتح وحدات النظام.",
        "منطقة المحتوى (مركز الشاشة): تَعرض لوحة القيادة التنفيذية افتراضياً.",
    ], 1):
        add_para(doc, f"{i}. {t}")
    add_image(doc, os.path.join(ASSETS, "01-dashboard.png"), caption="الشاشة الرئيسية ولوحة القيادة التنفيذية")
    add_h(doc, "٢-١  الأقسام التسعة في الشريط الجانبي", 3, color=TEAL)
    table = doc.add_table(rows=10, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["الجمهور المُستفيد", "الوحدات الفرعية", "القسم"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        set_para_rtl(p)
        r = p.add_run(h)
        r.bold = True
        set_run_arabic(r)
        r.font.size = Pt(11)
    rows = [
        ("كل المستخدمين", "الصفحة الرئيسية، المستفيدون", "الرئيسية"),
        ("الكوادر السريرية", "الملف الطبي، المتابعة اليومية، الأدوية", "الخدمات الطبية"),
        ("الاختصاصيون الاجتماعيون", "ملف الكرامة، محرك التمكين، بوابة الأسرة، البحث الاجتماعي", "الخدمات الاجتماعية"),
        ("قسم الجودة والإدارة العامة", "الحوكمة، المخاطر، الجودة، التميز، الدرع القانوني", "الحوكمة والجودة"),
        ("المالية والتشغيل", "لوحة التشغيل، الإعاشة، الأصول، المخزون", "العمليات"),
        ("المدير وقيادات الإدارة", "نبض المركز، التنبيهات الذكية، المكتبة الرقمية", "الذكاء والتنبؤ"),
        ("الإدارة العامة والوزارة", "لوحة التقارير، تَقرير العائد الاجتماعي", "التقارير"),
        ("المدراء والوكلاء", "بُوصلة القيادة", "القيادة الاستراتيجيّة"),
        ("المدير وإدارة الموارد البشرية", "الهيكل التنظيمي، الموظفون، الصلاحيات", "الإدارة"),
    ]
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            set_para_rtl(p)
            r = p.add_run(val)
            set_run_arabic(r)
            r.font.size = Pt(10)
    doc.add_paragraph()
    add_h(doc, "٢-٢  مَا تَعرضه لوحة القيادة فوراً", 3, color=TEAL)
    for b in [
        "بطاقات التنبيهات والمساءلة (عَدد القضايا المستحقة النَظَر، مع مستويات أهميَّة).",
        "مؤشِّرات الإنجاز: نسبة تَغطية الخطط، نسبة تَحقيق الأهداف، نسبة الأنشطة، نسبة المتابعات اليومية.",
        "عدد المستفيدين النشطين الإجمالي وفي كل عَنبَر.",
        "نَبض المركز (Vital Pulse) — مؤشِّر لحظي على حالة المخاطر العامة.",
        "مؤشِّر إحسان: مؤشِّر تَقدُّم نَوعي على جودة الخدمة الإنسانية.",
        "أداء الأقسام: لوحة بمعدَّل إنجاز الأهداف لكل قسم بصرياً.",
        "المهام المعلَّقَة القريبة من المستخدم.",
        "سجل الأنشطة الأخيرة.",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "للقارئ من القيادات",
                "اللوحة تُغني عن الطلب الدوري لتقارير 'كم مستفيداً نَخدم' أو 'ما مستوى الأداء' — الأرقام أمامك، مُحَدَّثَة لحظياً، قابلة للاستفسار العميق بنقرة.",
                role_color=NAVY)
    add_callout(doc, "للقارئ من مدير المركز",
                "الصورة العامة لمركزك في شاشة واحدة. كل بطاقة قابلة للنقر لمعرفة التَفاصيل ومَن يَتولى المعالجة.",
                role_color=ORANGE)
    add_callout(doc, "للقارئ من رؤساء الأقسام والوحدات",
                "بطاقات قسمك ضمن اللوحة تُمَيَّزُ لك تلقائياً، فتراها قبل غيرها.",
                role_color=TEAL)
    page_break(doc)

    # ---- SECTION 03 ----
    add_divider(doc, "٠٣", "المستفيدون — استعراض البيانات وملف الكرامة")
    add_h(doc, "٠٣  •  المستفيدون — استعراض البيانات وملف الكرامة", 2, color=NAVY)
    add_para(doc, "ركن النظام كلُّه. كل المُحَركات والوحدات تَعود في النهاية إلى ملف المستفيد.", size=12)
    add_h(doc, "٣-١  شاشة قائمة المستفيدين", 3, color=TEAL)
    add_para(doc, "عند النقر على 'المستفيدون' في الشريط الجانبي، تَنفتح قائمة المستفيدين النشطين في المركز.")
    add_image(doc, os.path.join(ASSETS, "basira-02-beneficiaries.png"), caption="شاشة قائمة المستفيدين")
    add_h(doc, "الميزات", 4, color=NAVY)
    for b in [
        "البحث المتقدِّم: بالاسم، الرقم، نوع الإعاقة، العنبر، الحالة.",
        "الفلترة الذكية: حالات الخطر العالي، حالات تَحتاج متابعة، قَوائم انتظار، أسماء قَيد التَخطيط.",
        "العرض المتعدِّد: قائمة، بطاقات، حسب العنبر.",
        "الإجراءات السريعة: فتح الملف، إضافة ملاحظة، تَسجيل حدث، تَكليف موظف.",
        "حالة الملف: مكتمل / مَنقوص / يَستلزم تَحديثاً.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "٣-٢  ملف المستفيد المتكامل", 3, color=TEAL)
    add_para(doc, "النَقر على اسم المستفيد يَفتح ملفه الكامل، ويَتألَّف من خمسة أقسام رئيسية:")
    for i, t in enumerate([
        "البَطاقة التَعريفية (الاسم، الرقم، تاريخ الالتحاق، العنبر، الفريق المُعالِج).",
        "الملف الطبي (التشخيصات الموثَّقَة، الأدوية، الحساسيات، تاريخ الحوادث).",
        "ملف الكرامة (مفصَّل أدناه).",
        "خطّة التأهيل والتمكين (الأهداف التأهيلية، نسب الإنجاز، التَدخُّلات).",
        "المتابعة اليومية (سجل الأنشطة، الاستهلاك الغذائي، النوم، الحالة المزاجية).",
    ], 1):
        add_para(doc, f"{i}. {t}")
    add_image(doc, os.path.join(ASSETS, "basira-16-beneficiary-detail.png"), caption="تفاصيل ملف المستفيد")
    add_h(doc, "٣-٣  ملف الكرامة — الميزة الجوهرية", 3, color=TEAL)
    add_para(doc, "ملف الكرامة هو ما يُمَيِّز بصيرة عن أيِّ نظام إدارة مَراكز سابق. يَحتوي:", bold=True)
    add_h(doc, "نوع الشخصية وأسلوب التواصل", 4, color=NAVY)
    for b in [
        "نوع الشخصية: اجتماعي / حيوي / هادئ / مُلاحظ.",
        "وَصف نَصيّ حُرّ.",
        "أسلوب التواصل المفضَّل: لفظي / بلغة الإشارة / بالإشارات / بالصور / مَختلَط.",
        "'أفضل أسلوب لإشراكه' — حقل نَصّي مفتوح، يَكتب فيه الموظف العارف به ما لا تَلتقطه الخانات.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "التَفضيلات الحسيَّة", 4, color=NAVY)
    for b in [
        "الإضاءة المفضَّلة (مُعتمَة / ساطعة / طبيعية / لا فرق).",
        "الضَجيج المُسْتساغ (هادئ / متوسط / حيوي).",
        "الحرارة (مُعتدل / دافئ / بارد).",
        "الروائح المُستحَبَّة أو المُنفِّرة.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "المحبوبات والمنفِّرات", 4, color=NAVY)
    for b in [
        "الطعام المفضَّل / المنفِّر.",
        "الأنشطة المحبَّبَة.",
        "الأماكن، الأشخاص، الألوان.",
        "المحفِّزات السلبية والمَخاوف.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "سجل الأعمال الكريمة (deeds)", 4, color=NAVY)
    add_para(doc, "ركنٌ مستقلّ يُوَثِّق:")
    for b in [
        "الأعمال الروحية (صلاة، ذكر، صيام، صدقة).",
        "الأعمال الاجتماعية (مساعدة، مشاركة، اعتذار، مصالحة).",
        "الأعمال الشخصية (إنجاز شخصي، تَجاوُز خوف).",
        "الأعمال الإبداعية (رَسم، صوت، حَرَكة، حِرَفَة).",
        "نَصرة الآخرين.",
    ]:
        add_bullet(doc, b)
    add_para(doc, "كل عمل يُوَثَّق مع: التاريخ، الفئة، مستوى الأثر، من شَهد عليه.")
    add_callout(doc, "القيمة التشغيلية لملف الكرامة",
                "هذا الملف هو أوَّل ما يَفتحه أي موظف يَلتقي بالمستفيد. ليس آخر ما يَفتح. الملف الطبي يَأتي بعد الكرامة — لا قبلها — لأن الكرامة هي سياق الحالة، لا تَفصيل لاحق منها.",
                role_color=GREEN)
    add_callout(doc, "للقارئ من رؤساء الأقسام",
                "ملف الكرامة مَطلوب من كل قسم تَحديثه ربع سنوياً، ضمن مؤشِّرات الأداء التي يَحسبها النظام تلقائياً.",
                role_color=TEAL)
    add_callout(doc, "للقارئ من إدارة الإشراف الاجتماعي",
                "جودة ملف الكرامة في كل مركز هي مؤشِّر مباشر على جودة معرفة المركز بمستفيديه. مَركز فيه ٧٠٪ من الملفَّات بحقول فارغة، يَستحقّ نَظرة إشرافية.",
                role_color=NAVY)
    page_break(doc)

    # ---- SECTION 04 ----
    add_divider(doc, "٠٤", "الخدمات الطبية والمتابعة اليومية")
    add_h(doc, "٠٤  •  الخدمات الطبية والمتابعة اليومية", 2, color=NAVY)
    add_para(doc, (
        "تُقدِّم وحدة الخدمات الطبية إطاراً متكامِلاً للجوانب الصحية ضمن خدمة التأهيل، دون أن تَحوِّل المركز إلى مستشفى. "
        "الخدمة الطبية في بصيرة مساعِدة للتأهيل الاجتماعي، لا أصلية."
    ), size=12)
    add_h(doc, "٤-١  تَقييم مَخاطر السقوط", 3, color=TEAL)
    add_para(doc, "شاشة تَخصصُّيَّة لرصد وتَقييم مخاطر السقوط، بناءً على معايير دولية مكيَّفَة لبيئة المركز.")
    add_image(doc, os.path.join(ASSETS, "basira-17-fall-risk.png"), caption="تَقييم مخاطر السقوط")
    add_h(doc, "آلية العمل", 4, color=NAVY)
    for i, t in enumerate([
        "الفريق يُجري تَقييماً لكل مستفيد عند الالتحاق، ثم دورياً.",
        "النَّظام يُولِّد تَلقائياً درجة خطورة وتَوصيات وقائية.",
        "التوصيات تَنفذ كمهام موَزَّعة على الكوادر.",
        "أي حادث يُسَجَّل ويَدخل في حساب الدرجة لاحقاً.",
    ], 1):
        add_para(doc, f"{i}. {t}")
    add_h(doc, "٤-٢  المتابعة اليومية", 3, color=TEAL)
    add_para(doc, "سجل يومي لكل مستفيد، يَضمّ:")
    for b in [
        "العلامات الحيوية (في حالات الإقامة الإيوائية): ضغط، نبض، حرارة، تَشبُّع الأكسجين، السكَّر للحالات المعنيَّة.",
        "استهلاك الوجبات والسوائل.",
        "حالة النوم والمزاج.",
        "الأنشطة المُنجَزَة في اليوم.",
        "ملاحظات الفريق (طبية، تَأهيلية، اجتماعية).",
    ]:
        add_bullet(doc, b)
    add_h(doc, "الميِّزات الذكية", 4, color=NAVY)
    for b in [
        "التَنبيه التلقائي: عند تَجاوز قراءة العلامات الحيوية للنطاق الطبيعي.",
        "التَتبُّع التَدريجي: رسم بياني لتَطوُّر العلامات على مدى أسابيع.",
        "ربط بالحوادث: تَغيُّر مفاجئ في مزاج المستفيد قد يَنعكس في تَنبيه إنذار مبكِّر.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "٤-٣  إدارة الأدوية", 3, color=TEAL)
    for b in [
        "جدول الأدوية الفعَّال لكل مستفيد.",
        "تَنبيه آلي بمواعيد الإعطاء.",
        "تَوثيق الإعطاء بِبصمَة الكادر المُسَجِّل.",
        "متابعة المخزون مع تَنبيه عند نَفاد الكميات.",
        "رَصد التَفاعُلات الدوائية المُحتمَلة.",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "للقارئ من القيادات السريرية",
                "المتابعة اليومية ليست فقط للتسجيل — هي مَصدر بيانات لمحرك حسّة زكية الذي يَسبق الأزمات. إذا كانت العلامات الحيوية لمستفيد تَتدهور تَدريجياً، يَكون التنبيه قَبل الأزمة بأيام، لا أثناءها.",
                role_color=NAVY)
    add_callout(doc, "للقارئ من رؤساء وحدات التَمريض",
                "كل مَناوبة تَستلم سجلاً واضحاً عن المستفيد، يَشمل ما حَدَث في النَوبة السابقة، بدون فقدان معلومة بسبب التَسليم.",
                role_color=TEAL)
    page_break(doc)

    # ---- SECTION 05 ----
    add_divider(doc, "٠٥", "الخدمات الاجتماعية ومحرك التمكين")
    add_h(doc, "٠٥  •  الخدمات الاجتماعية ومحرك التمكين", 2, color=NAVY)
    add_para(doc, "هنا تَكمن رسالة بصيرة الجَوهرية. إذا كان الملف الطبي ضرورياً، فالخَدَمات الاجتماعية ومحرك التَمكين هما سَبب وجود المركز.", size=12)
    add_h(doc, "٥-١  محرك التمكين", 3, color=TEAL)
    add_para(doc, "شاشة موحَّدة لكل خطط التأهيل والتَمكين الفردية.")
    add_image(doc, os.path.join(ASSETS, "basira-04-empowerment.png"), caption="محرك التمكين")
    add_h(doc, "مكوِّنات الخطّة", 4, color=NAVY)
    for i, t in enumerate([
        "الأهداف التأهيلية الفردية مَصاغَة بِمَنطق SMART.",
        "التَدخُّلات المخطَّطَة من قِبَل أي تَخصُّص.",
        "مؤشِّرات النَجاح: ما هي العلامات التي تَدلّ على التَقدُّم؟",
        "الجدول الزمني: مراجعات شهرية، فصلية، سنوية.",
        "شَهادات التَقدُّم: تَوثيق إنجازات يَستلمها المستفيد وأسرته.",
    ], 1):
        add_para(doc, f"{i}. {t}")
    add_h(doc, "التَوصيف الفلسفي للأهداف", 4, color=NAVY)
    add_para(doc, "كل هدف يجب أن يَكون مَوجَّهاً للمستفيد، لا للموظف:", bold=True)
    add_para(doc, "❌ 'تَطبيق برنامج العلاج الطبيعي ٣ مرات أسبوعياً' (هذا للموظف).", color="DC2626")
    add_para(doc, "✅ 'تَمكين المستفيد من النُهوض من السرير دون مساعدة، خلال ٣ أشهر' (هذا للمستفيد).", color="2BB574")
    add_h(doc, "٥-٢  بَوابة الأسرة", 3, color=TEAL)
    add_para(doc, "نَافذة الأسرة على المركز. تَفتح للأسرة (أو ولي الأمر) إمكانية:")
    for b in [
        "مُتابَعة الخطّة التَأهيلية للمستفيد بشكل دوري.",
        "رؤية الأنشطة اليومية والصور (بإذن الإدارة).",
        "التواصل مع الفريق المعالِج عبر قنوات مَحكَّمة.",
        "المُشاركة في تَخطيط الزيارات وإجازات نهاية الأسبوع.",
        "تَقديم الشكاوى والاقتراحات بسرّية.",
        "تَقييم رضاهم عبر استبيان شهري.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "الالتزامات الإطارية", 4, color=NAVY)
    for b in [
        "سرّية مَحكَمَة: لا يَستطيع أحد من الأسر رؤية بيانات أسرة أخرى.",
        "استجابة مَوقَّتة: ٢٤ ساعة للرسائل العامة، ٤ ساعات للأسئلة العاجلة، فوري للحالات الحرجة.",
        "حقوق المستفيد البالغ: المستفيد البالغ القادر على التعبير له القَول الفصل في ما يُكشف لأسرته من بياناته الحساسة.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "٥-٣  البحث الاجتماعي", 3, color=TEAL)
    add_para(doc, "الوحدة المختصَّة بالاختصاصيِّين الاجتماعيِّين:")
    for b in [
        "التَقييم الاجتماعي الأوَّلي عند الالتحاق.",
        "البحث الاجتماعي الدَوري (سنوياً أو عند الحاجة).",
        "خطط الإرشاد الأسري.",
        "مُتابَعَة الزِّيارات والإجازات.",
        "التَقييمات النَفسية الاجتماعية.",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "للقارئ من رؤساء الأقسام الاجتماعية",
                "كل تَقييم اجتماعي مُسَجَّل في النظام يُغذِّي ملف المستفيد ولا يَنسى. لا تَكرارَ في الأسئلة، لا فقدان لمعلومة بين سَنة وأخرى.",
                role_color=TEAL)
    add_callout(doc, "للقارئ من القيادات",
                "السطر الأخلاقي للنظام يَنعكس مباشرةً في هذه الوحدة. كل تَدَخُّل يُجيب عن سؤال 'أيّ حاجز اجتماعي نُفكِّكه؟'. إن غاب الجواب، التَدَخُّل يَستحقّ المراجعة.",
                role_color=NAVY)
    page_break(doc)

    # ---- SECTION 06 ----
    add_divider(doc, "٠٦", "الحوكمة والجودة والتميز المؤسسي")
    add_h(doc, "٠٦  •  الحوكمة والجودة والتميز المؤسسي", 2, color=NAVY)
    add_para(doc, "الجودة في بصيرة ليست مَلَفَّاً منفصلاً. هي طبقة عابرة تُوازِنُ كل وحدة. لكن لها مَركز: قسم الحوكمة والجودة في الشريط الجانبي.", size=12)
    add_h(doc, "٦-١  مَركز التميز المؤسسي", 3, color=TEAL)
    add_para(doc, "شاشة موحَّدة تَجمع كل أدوات الجودة والتميز.")
    add_image(doc, os.path.join(ASSETS, "basira-05-quality-excellence.png"), caption="مَركز التميز المؤسسي")
    add_h(doc, "الوحدات الفرعية", 4, color=NAVY)
    for b in [
        "الخيط الذهبي للحَوكَمَة: إطار مَوحَّد لتَتبُّع كل سياسة من اعتمادها إلى تَطبيقها.",
        "سجل المخاطر: مَنهجية مَوثَّقَة (ISO 31000) للتعامل مع المَخاطر.",
        "درع السلامة (مكافحة العدوى): ضَمان بيئة آمنة من الناحية الميكروبية.",
        "الامتثال ISO: تَوثيق الالتزام بـ ISO 9001:2015.",
        "دليل الجودة: المستندات الرسمية والإجراءات.",
        "سجل عدم المطابقة (NCR) والإجراءات التصحيحية والوقائية (CAPA).",
        "التدقيق الداخلي: دورات تَدقيق منتظمة بسجلات قابلة للمراجعة.",
        "لوحة الجودة: مؤشرات أداء الجودة في لوحة قيادة منفصلة.",
        "الدَرع القانوني: تَتبُّع الالتزام التَشريعي والقانوني.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "٦-٢  السلامة وَمكافَحَة العدوى (IPC)", 3, color=TEAL)
    add_para(doc, "شاشة تَخصُّصية لإدارة بيئة آمنة من الناحية الميكروبية.")
    add_image(doc, os.path.join(ASSETS, "basira-10-ipc.png"), caption="وحدة مكافحة العدوى")
    for b in [
        "التَفتيش اليومي: نَموذج تَفتيش مَنتظم لكل العنابر والمَرافق.",
        "بَلاغ الحوادث الميكروبية: نَموذج إبلاغ سريع.",
        "التَوعية والتَطعيمات: تَتبُّع تَطعيمات الكوادر والمستفيدين.",
        "مُعَدَّات الوقاية الشخصية (PPE): تَتبُّع المخزون والاستهلاك.",
        "إدارة تَفشِّي الحالات: سيناريوهات استجابة لأي تَفَشٍّ مُحتمَل.",
    ]:
        add_bullet(doc, b)
    add_para(doc, (
        "مهم: بصيرة لا يَعتمد على معايير CBAHI. مَراكز التَأهيل ليست مستشفيات. الإطار المَتَّبَع هو ISO 9001 (إدارة جودة عامة) "
        "وISO 45001 (سَلامة وصحة مهنية)، مع EQUASS كمَرجعية تَخصُّصية حيث تَنطبق."
    ), bold=True, color=NAVY)
    add_callout(doc, "للقارئ من قسم الجودة في كل مركز",
                "هذه الوحدة تَستبدل كل المَلَفَّات الورقية المتفرِّقة بنظام موحَّد لا يَفقد ذاكرة مؤسسية بِخُروج موظف.",
                role_color=TEAL)
    add_callout(doc, "للقارئ من إدارة الإشراف الاجتماعي",
                "جودة ما يُنفِّذه كل مركز مَكشوفة لك في لوحة موحَّدة. لا تَنتظر تَقريراً سنوياً متأخِّراً.",
                role_color=NAVY)
    add_callout(doc, "للقارئ من القيادات",
                "تَقارير الالتزام الجاهزة في النظام تُقدَّم مباشرة في تَقرير الوزارة السَنَوي عن أداء قطاع التنمية الاجتماعية، دون جَهد إضافي.",
                role_color=ORANGE)
    page_break(doc)

    # ---- SECTION 07 ----
    add_divider(doc, "٠٧", "العمليات اليومية")
    add_h(doc, "٠٧  •  العمليات اليومية", 2, color=NAVY)
    add_para(doc, "مَركز التأهيل الإيوائي مؤسسة تشغيلية كاملة: غذاء، صيانة، أصول، مَخزون، أمن. كل ذلك مُنَظَّم في وحدة العمليات.", size=12)
    add_h(doc, "٧-١  الإعاشة (وحدة التغذية)", 3, color=TEAL)
    add_para(doc, "شاشة متكامِلَة لإدارة وَجبات المركز اليومية.")
    add_image(doc, os.path.join(ASSETS, "basira-09-catering.png"), caption="وحدة الإعاشة")
    add_h(doc, "الوحدات الفرعية", 4, color=NAVY)
    for b in [
        "اللوحة العامة للإعاشة: نَظَرة على وَجبات اليوم، الكميات، المتلقِّين.",
        "السجل اليومي للإعاشة: تَوثيق ما قُدِّم فعلاً، ومَن استَهلك ومَن لا.",
        "الفاتورة الشهرية: حساب آلي لكلفة الإعاشة الشهرية.",
        "مُراقبة الجودة: نَماذج فَحص جودة الوجبات الواردة.",
        "لوحة جودة الإعاشة: مؤشِّرات صحَّيَّة وغذائية.",
        "تَقارير الإعاشة: تَقارير دَوريَّة (يومية، أسبوعية، شهرية).",
    ]:
        add_bullet(doc, b)
    add_h(doc, "الميزات الذكية", 4, color=NAVY)
    for b in [
        "التَكامل مع الأنظمة الغذائية الفردية.",
        "حساب القيمة الغذائية: تَوازن العناصر في كل وجبة.",
        "إدارة المُورِّدين: تَقييم جَودة المُورِّد، تَوثيق الفواتير.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "٧-٢  الأصول والصيانة", 3, color=TEAL)
    for b in [
        "سجل الأصول: كل أصل في المركز مُسَجَّل.",
        "جَدول الصيانة الوقائية: مواعيد دَورية لصيانة كل أصل.",
        "بَلاغ الأعطال: نَموذج إبلاغ سريع.",
        "تاريخ الصيانة: لكل أصل تاريخ كامل من إصلاحات وتَكاليف.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "٧-٣  المخزون والكسوة", 3, color=TEAL)
    for b in [
        "إدارة مخزون الكسوة: لكل مستفيد قائمة احتياجات وَتَوزيع.",
        "إدارة المخزون العام: مَواد التشغيل، المعدات الاستهلاكية.",
        "تَنبيه عند نَفاد المخزون: لا تُستنزَف مادة فجأة.",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "للقارئ من المدير الإداري",
                "هذه الوحدة تُغني عن خمسة دفاتر وَرَقية على الأقلّ. كل العمليَّات الإدارية مَوثَّقَة وقابلة للتَدقيق.",
                role_color=TEAL)
    add_callout(doc, "للقارئ من القيادات",
                "التَكاليف التَشغيلية الفعلية مَحسوبة آلياً، ومَتاحة لتقارير الكفاءة الإنفاقية.",
                role_color=NAVY)
    page_break(doc)

    # ---- SECTION 08 ----
    add_divider(doc, "٠٨", "الذكاء والتنبؤ — الإنذار المبكِّر ونبض المركز")
    add_h(doc, "٠٨  •  الذكاء والتنبؤ", 2, color=NAVY)
    add_para(doc, "ركن الفَرق النَوعي في بصيرة. هنا تَتلاقى البيانات والذكاء المؤسَّسي ليُنبِّها قبل الأزمة، لا أثناءها.", size=12)
    add_h(doc, "٨-١  مَركز المؤشِّرات الذكية", 3, color=TEAL)
    add_para(doc, "شاشة موحَّدة لكل مؤشِّرات الأداء.")
    add_image(doc, os.path.join(ASSETS, "basira-06-indicators.png"), caption="مَركز المؤشِّرات الذكية")
    add_h(doc, "المؤشِّرات الموجودة", 4, color=NAVY)
    for b in [
        "مؤشِّرات الإنذار المبكِّر: حالات تَستحقّ الانتباه قبل أن تَتدهور.",
        "التنبُّؤ السلوكي: تَوقُّعات على سلوك المستفيد بناءً على تَاريخه.",
        "التَدقيق البيولوجي: تَتبُّع العلامات الحيوية اللازمة.",
        "مؤشرات الرضا (نَبض الرضا): رضا المستفيد، رضا الأسرة، رضا الكوادر.",
        "مؤشرات الكلفة لكل مستفيد.",
        "مؤشِّرات الأثر على الموارد البشرية (HR Impact).",
        "الالتزام بـ ISO.",
        "المعدَّلات المعيارية (Benchmarks).",
        "مؤشرات الأهداف الاستراتيجية.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "٨-٢  نظام الإنذار المبكِّر", 3, color=TEAL)
    add_para(doc, "شاشة تَخصُّصية للحالات التي تَستحقّ التَدخُّل قبل تَفاقُمها.")
    add_image(doc, os.path.join(ASSETS, "basira-12-early-warning.png"), caption="نظام الإنذار المبكِّر")
    add_h(doc, "آلية العمل", 4, color=NAVY)
    for i, t in enumerate([
        "النَّظام يَجمع بيانات يومية من كل المصادر.",
        "محرك حسّة زكية يُحلِّل التراكُمات.",
        "عند تَجاوُز عَتَبَة معيَّنة، يَصدر تَنبيه ينتقل إلى المسؤول المعنيّ.",
        "التَنبيه يَكون شفّافاً قابلاً للتفسير: لماذا صَدَر، ما العوامل، ما المُقتَرَح.",
        "القرار يَبقى للإنسان دائماً. النَّظام يُنبِّه، ولا يُقَرِّر.",
    ], 1):
        add_para(doc, f"{i}. {t}")
    add_h(doc, "أنواع الإنذارات", 4, color=NAVY)
    for b in [
        "إنذار سريري: تَدهور مَلحوظ في علامة حيوية.",
        "إنذار سلوكي: تَغيُّر مفاجئ في نمط النَوم أو المزاج أو التَفاعُل.",
        "إنذار اجتماعي: عُزلَة متزايدة، انقطاع تَواصُل أُسري.",
        "إنذار تَأهيلي: تَوقُّف التَقدُّم في خطّة مستفيد.",
        "إنذار جماعي (على مستوى العنبر): عَنبَر يَزداد فيه ضغط حالات الخطر العالي.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "٨-٣  نَبض الرضا (Satisfaction Pulse)", 3, color=TEAL)
    add_para(doc, "شاشة لرَصد رضا المستفيد وذويه.")
    add_image(doc, os.path.join(ASSETS, "basira-14-satisfaction.png"), caption="نَبض الرضا")
    add_h(doc, "كَيف يُقاس", 4, color=NAVY)
    for b in [
        "استبيان شهري للأسرة عبر بَوابة الأسرة.",
        "تَقييم نَوعي ربع سنوي: حوار مباشر مع الأسرة، يُسَجَّل نَصيَّاً، ويُحلَّل دلالياً.",
        "صوت المستفيد المباشر (لمن يَستطيع التَعبير): تَسجيل صوتي / نَصي / رَمزي شَهري.",
    ]:
        add_bullet(doc, b)
    add_para(doc, "التزام جوهري: صوت المستفيد لا يُعاقَب عليه أحد. هو مَدخَل تَحسين، لا أداة عقاب.", bold=True, color=GREEN)
    add_h(doc, "٨-٤  المكتبة الرقمية", 3, color=TEAL)
    add_para(doc, "شاشة المرجعيَّات والمعارف.")
    add_image(doc, os.path.join(ASSETS, "basira-15-knowledge.png"), caption="المكتبة الرقمية")
    add_callout(doc, "للقارئ من القيادات",
                "ما يُولِّده النظام من مؤشِّرات يُغني عن طلب التقارير المتفرِّقة. صورة المركز في لوحة واحدة، تَفاصيلها قابلة للاستكشاف.",
                role_color=NAVY)
    add_callout(doc, "للقارئ من رؤساء الأقسام",
                "الإنذار المبكِّر يُتيح لك التَدخُّل وَقت يَصنع فرقاً، لا بعد فَوات الأوان.",
                role_color=TEAL)
    add_callout(doc, "للقارئ من إدارة الإشراف الاجتماعي",
                "المؤشِّرات الذكية تَنقل التَدقيق من الزيارات السنوية إلى المراقبة الذكية المستمرَّة، مع الزيارات للحالات التي تَستحقّ.",
                role_color=ORANGE)
    page_break(doc)

    # ---- SECTION 09 ----
    add_divider(doc, "٠٩", "التقارير والمؤشرات الاستراتيجية")
    add_h(doc, "٠٩  •  التقارير والمؤشرات الاستراتيجية", 2, color=NAVY)
    add_para(doc, "البَيانات لا تَكتسب قيمتها إلا حين تُتَرجَم إلى تَقارير قابلة للقراءة، تُعِين القيادة على القرار. وحدة التقارير هي ذلك الجسر.", size=12)
    add_h(doc, "٩-١  التقرير التنفيذي", 3, color=TEAL)
    add_para(doc, "شاشة موَحَّدة لتَوليد التقارير التَنفيذية.")
    add_image(doc, os.path.join(ASSETS, "basira-07-executive-report.png"), caption="التقرير التنفيذي")
    add_h(doc, "أنواع التقارير", 4, color=NAVY)
    for b in [
        "تَقرير المركز الشهري: نَظَرة عامة على الأداء، المؤشِّرات، الأحداث، الإنجازات.",
        "تَقرير تَحليلي: تَفصيل أكثر عمقاً، مَع مقارنات زمنية.",
        "تَقرير مُخصَّص: ينشئ المستخدم التَقرير بحَسب احتياجه.",
        "تَقرير العائد الاجتماعي للاستثمار (SROI): حساب القيمة الاجتماعية المُولَّدة لكل ريال مَصروف.",
        "تقارير تَستوفي متطلبات حقيبة مؤشِّرات الأداء الوزارية.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "٩-٢  المؤشِّرات الاستراتيجية", 3, color=TEAL)
    add_para(doc, "شاشة مَوصولة برؤية الوزارة الاستراتيجية ورؤية ٢٠٣٠.")
    add_image(doc, os.path.join(ASSETS, "basira-13-strategic-kpi.png"), caption="المؤشِّرات الاستراتيجية")
    for b in [
        "المُسْتهدَفات الاستراتيجية للمركز ومستوى تَحقيقها.",
        "المُسْتهدَفات على مستوى المنطقة.",
        "المُسْتهدَفات الوزارية المرتبطة بالمركز.",
        "محاور رؤية ٢٠٣٠ (جودة الحياة، تَنمية القدرات البشرية).",
        "اتفاقية CRPD ومُؤَشِّرات الالتزام بمَوادها.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "٩-٣  تَقرير العائد الاجتماعي SROI", 3, color=TEAL)
    add_para(doc, "من أبرز مَيزات بصيرة الاستراتيجية. لكل ريال مَصروف على المركز، ما هو مَقدار القيمة الاجتماعية المُولَّدَة؟", bold=True)
    add_h(doc, "كيف يُحَسب", 4, color=NAVY)
    for b in [
        "التَكاليف المباشرة: إيواء، رواتب، أدوية، إعاشة، أنشطة.",
        "القيمة المُولَّدة: تَخفيض احتياج الأسرة لرعاية مدفوعة، فرص العمل المُستحصَلَة، تَحسُّن الاستقلالية.",
        "القيمة النَوعية: الكرامة، حقّ المشاركة، أَثَر الأسرة (تُذكَر نصيَّاً لا رقمياً).",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "للقارئ من القيادات العليا والإدارة العامة",
                "التَقرير التَنفيذي يُجَهِّز لك الإجابة على سؤال 'ما الذي تُحدِثه ميزانية مَراكز التأهيل؟'، بأرقام موثَّقَة، لا انطباعات.",
                role_color=NAVY)
    add_callout(doc, "للقارئ من المدير العام بفرع الوزارة",
                "المؤشِّرات الاستراتيجية تُتيح لك مقارنة المراكز ضمن فرعك، لا للمحاسبة العقابية، بل لتَوجيه الدَّعم لمَن يَحتاجه.",
                role_color=ORANGE)
    add_callout(doc, "للقارئ من إدارة الإشراف الاجتماعي",
                "التَقرير المؤسسي يُغني عن جَولات التَفتيش المتكرِّرة لجَمع نَفس البيانات. التَفتيش يَتركَّز على ما لا تَكشفه البيانات.",
                role_color=TEAL)
    page_break(doc)

    # ---- SECTION 10 ----
    add_divider(doc, "١٠", "بُوصلة القيادة الاستراتيجية")
    add_h(doc, "١٠  •  بُوصلة القيادة الاستراتيجية", 2, color=NAVY)
    add_para(doc, (
        "من أكثر مَيزات بصيرة ابتكاراً. القرارات الاستراتيجية في المؤسسات الكبرى تُعاني من ثلاثة أمراض: "
        "تُتَّخذ بدون توثيق سياقها، تُنسى بمرور الزمن، ولا تُربَط بأثرها. بُوصلة القيادة تُعالج الأمراض الثلاثة."
    ), size=12)
    add_image(doc, os.path.join(ASSETS, "basira-08-leadership-compass.png"), caption="بوصلة القيادة الاستراتيجية")
    add_h(doc, "١٠-١  مكوِّنات البُوصلة", 3, color=TEAL)
    add_h(doc, "بطاقات القرار (Decision Cards)", 4, color=NAVY)
    add_para(doc, "لكل قَرار استراتيجي يُتَّخَذ، تُسجَّل بطاقة مُوَحَّدة:")
    for b in [
        "السياق: ما الذي اقتَضى القرار؟",
        "البدائل المطروحَة: ما الخَيارات التي وُضعَت على الطاولة؟",
        "البديل المختار: ما الذي اخترناه؟ ولماذا؟",
        "البدائل المرفوضَة: لماذا رُفِضَت كل واحدة منها بشكل محدَّد؟",
        "أصحاب المصلحة: مَن يَتأَثَّر؟",
        "الجدول الزمني للمراجعة: متى نَعود لهذا القرار؟",
    ]:
        add_bullet(doc, b)
    add_h(doc, "النَتائج المرآة (Mirror Findings)", 4, color=NAVY)
    add_para(doc, "ملاحظات داخلية تَكشفها البيانات (لا التَفتيش):")
    add_para(doc, "'العنبر أ تَكَرَّر فيه استخدام الكرسي المتحرِّك ٣٠٪ أقلّ من العنبر ب رغم تَشابه ملفَّات المستفيدين — لماذا؟'", italic=True, color=GRAY)
    add_para(doc, "'نسبة المستفيدين الذين تَحقَّقوا أهدافهم في الربع الأخير ارتَفَعَت ١٥٪ — أي تَدخُّل تَسبَّب في ذلك؟'", italic=True, color=GRAY)
    add_h(doc, "المسارات (Trajectories)", 4, color=NAVY)
    add_para(doc, "تَتبُّع زمني لمؤشِّر مُعَيَّن: هل يَتحسَّن أم يَتراجع؟ ما العوامل المرتبطة بالتغيُّر؟ لا يُتَّخَذ قَرار جديد إلا بعد مراجعة المسار.")
    add_h(doc, "أُفُق السياسة (Policy Horizon)", 4, color=NAVY)
    add_para(doc, "السياسات الوزارية القادمة المتوقَّعَة، وأَثَرها المُحتَمَل على المركز، حتى يَستعدَّ القائد لها قَبل أن تَفاجِئَه.")
    add_h(doc, "١٠-٢  القيمة المؤسسية", 3, color=TEAL)
    add_para(doc, "بدلاً من أن تُسلَّم القيادة الجديدة 'ملف قرارات سابقة' مكتوبة في إيميلات وملاحظات متفرِّقة، تُسلَّم بُوصلة قيادة تَنبض حياً.")
    for b in [
        "يَستلم المدير الجديد بُوصلة جاهزة فيها كل قرار سابق بسياقه.",
        "يَستطيع تَتبُّع لماذا اتُّخذ كل قرار.",
        "يَرى أثره الفعلي.",
        "يَقرِّر إن كان يَستحقّ المراجعة.",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "للقارئ من المدير العام بفرع الوزارة",
                "بُوصلة القيادة تَجمع لك قرارات كل المراكز في منطقتك، فترى الأنماط، تَكتشف الإخفاقات المتكرِّرة، وتُوَزِّع الدَّعم وفقها.",
                role_color=ORANGE)
    add_callout(doc, "للقارئ من القيادات العليا",
                "بُوصلة على مستوى الوكالة تَنشأ من تَجميع بُوصلات المراكز. تَرى نمط القرار في القطاع كله، ومَن يَتَّخذ قرارات جيِّدة.",
                role_color=NAVY)
    add_callout(doc, "للقارئ من مدير المركز",
                "كل قرار تُسَجِّله يُحفَظ بسياقه. لا يُتَّهم لاحقاً أحد بقرار اتُّخذَ في وقت ضيِّق بدون فهم ظَرفه.",
                role_color=TEAL)
    page_break(doc)

    # ---- SECTION 11 ----
    add_divider(doc, "١١", "الإدارة والصلاحيات وسجلات التدقيق")
    add_h(doc, "١١  •  الإدارة والصلاحيات وسجلات التدقيق", 2, color=NAVY)
    add_para(doc, "ركن النظام الإداري والتقني. هنا تُعَالَج الكوادر، الصلاحيات، والسجلات الكاملة لأنشطة المستخدمين.", size=12)
    add_h(doc, "١١-١  سجل التَدقيق", 3, color=TEAL)
    add_para(doc, "أكثر شَاشة احترافية في النظام، يَخضع لها كل تَفاعُل.")
    add_image(doc, os.path.join(ASSETS, "basira-11-admin-audit.png"), caption="سجل التَدقيق")
    add_h(doc, "ما يُسَجَّل", 4, color=NAVY)
    add_para(doc, "كل عمل في النظام مَوثَّق بـ:")
    for b in [
        "مُعرِّف المستخدم: مَن قام بالعمل.",
        "العمل: ما الذي حَدَث (قراءة، كتابة، تَعديل، حذف).",
        "الكائن: على ماذا (مَلَفّ مستفيد، تَنبيه، تَقييم، تَقرير).",
        "الطابع الزمني: متى بِدِقَّة الميلي ثانية.",
        "عنوان IP: من أيِّ جهاز.",
        "الجلسة: تَتبُّع كامل للجلسة.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "السجلات غير قابلة للتَعديل", 4, color=NAVY)
    for b in [
        "لا يَستطيع أي مستخدم — حتى المدير — تَعديل أو حذف سَجِلّ مَوجود.",
        "السجلات Append-Only.",
        "الاحتفاظ لِمدَّة ٧ سنوات.",
        "صَالحَة كمرجع قانوني عند الحاجة.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "١١-٢  تَطبيق التزامات الأمن السيبراني", 3, color=TEAL)
    add_para(doc, "السجلات تَخدم:")
    for b in [
        "سياسة DT-IS-POL-1300 V7 (إدارة سجلَّات الأحداث والمراقبة).",
        "نظام حماية البيانات الشخصية (PDPL).",
        "الضوابط الأساسية للأمن السيبراني NCA ECC-2:2024.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "١١-٣  الهيكل التَنظيمي والصَلاحيات", 3, color=TEAL)
    add_para(doc, "شاشة مرافِقَة لإدارة الكوادر:")
    for b in [
        "هيكل المركز: المدير، رؤساء الأقسام، الكوادر، تَوزيع العنابر.",
        "بَطاقات الكوادر: مع التَخصُّصات والشَهادات.",
        "تَخصيص الصَلاحيات: لكل دور صَلاحياته المحدَّدَة.",
        "التَدريب والتَطوير: تَتبُّع الدورات والشَهادات لكل موظف.",
        "تَقييم الأداء: نَموذج تَقييم سَنَوي مَوصول بمؤشِّرات الأداء الفعلية.",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "للقارئ من إدارة الموارد البشرية في المركز",
                "كل موظف لديه ملفّ كامل، يَتطوَّر معه. لا يَفقد إنجازَه إذا انتقل بين الأقسام.",
                role_color=TEAL)
    add_callout(doc, "للقارئ من إدارة الإشراف الاجتماعي",
                "التَدقيق يَستطيع كَشف أنماط مهمَّة (مثلاً: مَن يَدخل النظام في أوقات غير اعتيادية؟ مَن يُجري عمليَّات بكثافة على ملف معيَّن؟). أداة استكشافية، لا فقط أرشيفية.",
                role_color=NAVY)
    add_callout(doc, "للقارئ من القيادات في الوكالات المركزية",
                "التَدقيق على مستوى الوزارة يَستحقّ لجنة دَورية لمراجعته (يُنفَّذ في مَركز عمليَّات الأمن SOC). أنماط غير اعتيادية تَكشف ما تَفوته الزيارات الميدانية.",
                role_color=ORANGE)
    page_break(doc)

    # ---- SECTION 12 ----
    add_divider(doc, "١٢", "لوحة القيادة الإشرافية الإقليمية")
    add_h(doc, "١٢  •  لوحة القيادة الإشرافية الإقليمية", 2, color=NAVY)
    add_para(doc, "عند التَعميم على مَراكز متعدِّدة، يَنتقل النظام إلى مستوى تَجميعي يَخدم القيادات فوق المركز.", size=12)
    add_image(doc, os.path.join(ASSETS, "basira-18-aggregate.png"), caption="لوحة القيادة الإشرافية الإقليمية")
    add_h(doc, "١٢-١  ما تَعرضه اللوحة", 3, color=TEAL)
    for b in [
        "رؤية موحَّدة لكل المراكز ضمن الفرع/المنطقة/الوكالة.",
        "مقارنة شفّافة بين المراكز (لا للمحاسبة العقابية، بل لتَوزيع الدَّعم).",
        "رصد المخاطر مركزياً: مَركز تَتراكم فيه إشارات احتراق وظيفي / تَدنِّي رضا / تَأخُّر في تَحديث الملفات.",
        "المؤشِّرات المُقارَنَة: نسبة تَفعيل بَوابة الأسرة، تَحديث ملف الكرامة، الحوادث، تَحقيق الأهداف.",
        "التَنبيهات الإشرافية: حالات تَستحقّ زيارة ميدانية.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "١٢-٢  مَن يَستفيد من اللوحة", 3, color=TEAL)
    add_h(doc, "المدير العام بفرع الوزارة", 4, color=NAVY)
    for b in [
        "لوحة على مستوى الفرع كاملاً.",
        "مَركز واحد في الفرع متراجع → دَعم استثنائي بدلاً من عقوبة.",
        "مَركز متفوِّق → دراسة تَجربته لنَقلها للآخرين.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "إدارة الإشراف الاجتماعي بفرع الوزارة", 4, color=NAVY)
    for b in [
        "متابعة جودة الخدمات الاجتماعية في كل المراكز ضمن الفرع.",
        "التَأكُّد من تَطبيق سياسات الوزارة بشكل موَحَّد.",
        "تَدقيق مَلَفَّات الكرامة، تَحقيق الأهداف، رضا المستفيدين.",
        "تَوجيه الدَّعم الفنّي للمراكز التي تَحتاج دعماً.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "الإدارة العامة لخدمات الفروع", 4, color=NAVY)
    for b in [
        "لوحة على مستوى المملكة كاملاً.",
        "مقارنة المناطق.",
        "تَوزيع الموارد على أساس الاحتياج المُقاس.",
        "تَقارير لمعالي الوزير ولمَجلس الشؤون الاقتصادية والتَنمية.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "اللجنة العليا لبصيرة (مَجلس الإدارة)", 4, color=NAVY)
    for b in [
        "النَّظرة الاستراتيجية الشاملة.",
        "اعتماد التَطوير القادم.",
        "توجيه السياسات.",
    ]:
        add_bullet(doc, b)
    add_h(doc, "١٢-٣  المؤشِّرات الإشرافية الرئيسية", 3, color=TEAL)
    for b in [
        "نسبة تَفعيل بَوابة الأسرة في كل مَركز.",
        "نسبة تَحديث ملف الكرامة.",
        "متوسط الأهداف التَأهيلية المحَقَّقَة شهرياً لكل مستفيد.",
        "نسبة الموظفين تحت 'خط الاحتراق' (محرك مروءة).",
        "نسبة الحوادث.",
        "مؤشِّر رضا المستفيد.",
        "نسبة الالتزام بـ PDPL.",
        "توفُّر النظام (Uptime).",
    ]:
        add_bullet(doc, b)
    add_callout(doc, "للقارئ من المدير العام بفرع الوزارة",
                "لوحة الفَرع تُغني عن طلب تَقارير ربع سنوية متأخِّرة. الرؤية لحظية، الاستفسار العميق متاح بنقرة.",
                role_color=ORANGE)
    add_callout(doc, "للقارئ من إدارة الإشراف الاجتماعي",
                "المؤشِّرات تُحدِّد لك أين تَزور، ولماذا، ومَع مَن تَتحدَّث. الزيارة الميدانية تَكتسب قيمة أعلى لأنها مَوجَّهة.",
                role_color=NAVY)
    add_callout(doc, "للقارئ من القيادات العليا في الوكالات المركزية",
                "على مستوى الوزارة، اللوحة تُجيب: هل تَتقدَّم المملكة في تَفكيك الحواجز الاجتماعية لذوي الإعاقة؟ — بإحصاءات قابلة للنَشر في تَقارير الرؤية وتَقارير CRPD الدورية.",
                role_color=TEAL)
    page_break(doc)

    # ---- CLOSING ----
    add_h(doc, "خَاتمة الدليل", 1, color=NAVY, size=32)
    add_para(doc, (
        "نظام بصيرة ليس 'نظام إدارة' بالمعنى التقليدي. هو بنية تَفكير مؤسسي لخدمة الأشخاص ذوي الإعاقة، "
        "مُعَبَّرٌ عنه في تَطبيق رقمي. هذا الدليل قَدَّم خَريطة الواجهات والوظائف، لكن القيمة الفعلية تَنشأ في الاستخدام اليومي:"
    ), size=12)
    for b in [
        "موظف يَفتح ملف الكرامة قَبل الملف الطبي.",
        "مدير يَستخدم بُوصلة القيادة قبل اجتماع.",
        "اختصاصي اجتماعي يَكتب هدفاً للمستفيد، لا للموظف.",
        "أسرة تَطمَئنُّ على وَلَدها عبر بَوابة موثوقة.",
        "قيادة عُليا تَملك بياناتها لا انطباعاتها.",
    ]:
        add_bullet(doc, b)
    add_para(doc, "كل واجهة في هذا الدليل أُعِدَّت لتَصبح عادةً. والعادة، حين تَتراكم في مؤسسة كاملة، تَصير ثقافة.", bold=True)
    doc.add_paragraph()
    p_q = doc.add_paragraph()
    set_para_rtl(p_q)
    p_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rq = p_q.add_run(
        "البَصيرة ليست أن نَرى ما يَراه غيرنا، بل أن نُدرك ما خَفِيَ عنهم. ما يَخفى عن النظام التَقليدي هو الإنسان وراء الحالة. ما يَستحضره بصيرة هو ذلك الإنسان."
    )
    set_run_arabic(rq)
    rq.italic = True
    rq.font.size = Pt(13)
    rq.font.color.rgb = RGBColor.from_string(TEAL)
    page_break(doc)

    # ---- BACK COVER ----
    for _ in range(15):
        doc.add_paragraph()
    p_bc = doc.add_paragraph()
    p_bc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rbc = p_bc.add_run("www.hrsd.gov.sa")
    set_run_arabic(rbc)
    rbc.font.size = Pt(14)
    rbc.font.color.rgb = RGBColor.from_string(NAVY)

    # SAVE
    doc.save(OUT_DOCX)
    print(f"OK saved: {OUT_DOCX}")
    print(f"Size: {os.path.getsize(OUT_DOCX) / 1024:.1f} KB")


if __name__ == "__main__":
    main()
